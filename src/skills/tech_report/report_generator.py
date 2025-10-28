"""
Professional technical report generator.
Creates investor and client-ready documents from performance analysis.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, List, Optional, Any
import os


class ReportStyle:
    """Professional styling configuration for reports."""
    
    # Color scheme (matching performance analyzer)
    PRIMARY = RGBColor(46, 134, 171)      # Blue
    SECONDARY = RGBColor(162, 59, 114)    # Purple
    ACCENT = RGBColor(241, 143, 1)        # Orange
    SUCCESS = RGBColor(6, 167, 125)       # Green
    DANGER = RGBColor(199, 62, 29)        # Red
    DARK = RGBColor(43, 45, 66)          # Dark blue-gray
    LIGHT_GRAY = RGBColor(108, 117, 125) # Gray
    
    # Typography
    TITLE_FONT = 'Calibri'
    BODY_FONT = 'Calibri'
    MONO_FONT = 'Consolas'
    
    TITLE_SIZE = 26
    HEADING1_SIZE = 18
    HEADING2_SIZE = 14
    BODY_SIZE = 11
    CAPTION_SIZE = 9


class TechnicalReportGenerator:
    """Generates professional technical reports."""
    
    def __init__(self, report_type: str = 'comprehensive'):
        """
        Initialize report generator.
        
        Args:
            report_type: Type of report to generate
                - 'executive': High-level executive summary
                - 'technical': Detailed technical analysis
                - 'benchmark': Performance benchmark report
                - 'comprehensive': Full report with all sections
        """
        self.report_type = report_type
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure document styles."""
        styles = self.doc.styles
        
        # Heading 1
        try:
            h1 = styles['Heading 1']
            h1.font.name = ReportStyle.TITLE_FONT
            h1.font.size = Pt(ReportStyle.HEADING1_SIZE)
            h1.font.color.rgb = ReportStyle.PRIMARY
            h1.font.bold = True
        except KeyError:
            pass
        
        # Heading 2
        try:
            h2 = styles['Heading 2']
            h2.font.name = ReportStyle.TITLE_FONT
            h2.font.size = Pt(ReportStyle.HEADING2_SIZE)
            h2.font.color.rgb = ReportStyle.DARK
            h2.font.bold = True
        except KeyError:
            pass
        
        # Normal body text
        try:
            normal = styles['Normal']
            normal.font.name = ReportStyle.BODY_FONT
            normal.font.size = Pt(ReportStyle.BODY_SIZE)
        except KeyError:
            pass
    
    def add_cover_page(self, 
                      title: str,
                      subtitle: Optional[str] = None,
                      company: str = "SAAAM LLC",
                      date: Optional[datetime] = None):
        """
        Add professional cover page.
        
        Args:
            title: Report title
            subtitle: Optional subtitle
            company: Company name
            date: Report date (defaults to now)
        """
        if date is None:
            date = datetime.now()
        
        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(ReportStyle.TITLE_SIZE)
        title_run.font.color.rgb = ReportStyle.PRIMARY
        title_run.font.bold = True
        title_run.font.name = ReportStyle.TITLE_FONT
        
        # Subtitle
        if subtitle:
            self.doc.add_paragraph()  # Spacing
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(ReportStyle.HEADING2_SIZE)
            subtitle_run.font.color.rgb = ReportStyle.DARK
            subtitle_run.font.name = ReportStyle.TITLE_FONT
        
        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Company and date
        footer_para = self.doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"{company}\n{date.strftime('%B %d, %Y')}"
        )
        footer_run.font.size = Pt(ReportStyle.BODY_SIZE)
        footer_run.font.color.rgb = ReportStyle.LIGHT_GRAY
        
        # Page break
        self.doc.add_page_break()
    
    def add_executive_summary(self, 
                            summary: str,
                            key_findings: List[str],
                            metrics: Dict[str, Any]):
        """
        Add executive summary section (BLUF - Bottom Line Up Front).
        
        Args:
            summary: High-level summary text
            key_findings: List of key findings
            metrics: Key performance metrics to highlight
        """
        # Section heading
        self.doc.add_heading('Executive Summary', level=1)
        
        # Add highlight box with key metric
        if metrics:
            metric_para = self.doc.add_paragraph()
            metric_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Get most important metric
            if 'speedup' in metrics:
                metric_text = f"🚀 {metrics['speedup']:.1f}x Performance Improvement"
                color = ReportStyle.SUCCESS
            elif 'tflops' in metrics:
                metric_text = f"⚡ {metrics['tflops']:.2f} TFLOPS Achieved"
                color = ReportStyle.PRIMARY
            else:
                first_key = list(metrics.keys())[0]
                metric_text = f"{first_key}: {metrics[first_key]}"
                color = ReportStyle.PRIMARY
            
            metric_run = metric_para.add_run(metric_text)
            metric_run.font.size = Pt(16)
            metric_run.font.bold = True
            metric_run.font.color.rgb = color
            
            self.doc.add_paragraph()  # Spacing
        
        # Summary text
        summary_para = self.doc.add_paragraph(summary)
        summary_para.style = 'Normal'
        
        # Key findings
        if key_findings:
            self.doc.add_paragraph()
            findings_heading = self.doc.add_paragraph()
            findings_run = findings_heading.add_run('Key Findings:')
            findings_run.font.bold = True
            findings_run.font.size = Pt(ReportStyle.BODY_SIZE + 1)
            
            for finding in key_findings:
                bullet_para = self.doc.add_paragraph(finding, style='List Bullet')
                bullet_para.style = 'List Bullet'
    
    def add_performance_overview(self,
                                results: List[Dict[str, Any]],
                                baseline: Optional[Dict[str, Any]] = None):
        """
        Add performance overview section with metrics table.
        
        Args:
            results: List of performance results
            baseline: Optional baseline for comparison
        """
        self.doc.add_heading('Performance Overview', level=1)
        
        if not results:
            self.doc.add_paragraph("No performance data available.")
            return
        
        # Create table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Metric', 'Value', 'Unit', 'vs Baseline']
        for i, header in enumerate(headers):
            cell_para = header_cells[i].paragraphs[0]
            cell_run = cell_para.add_run(header)
            cell_run.font.bold = True
            cell_run.font.size = Pt(ReportStyle.BODY_SIZE)
        
        # Data rows
        for result in results:
            row_cells = table.add_row().cells
            
            # Metric name
            row_cells[0].text = result.get('name', 'Unknown')
            
            # Value
            value = result.get('value', 0)
            row_cells[1].text = f"{value:.2f}"
            
            # Unit
            row_cells[2].text = result.get('metric_type', '')
            
            # Comparison
            if baseline:
                baseline_val = baseline.get('value', 0)
                if baseline_val > 0:
                    change = ((value - baseline_val) / baseline_val) * 100
                    comparison_para = row_cells[3].paragraphs[0]
                    comparison_run = comparison_para.add_run(f"{change:+.1f}%")
                    
                    # Color code
                    if change > 5:
                        comparison_run.font.color.rgb = ReportStyle.SUCCESS
                    elif change < -5:
                        comparison_run.font.color.rgb = ReportStyle.DANGER
                else:
                    row_cells[3].text = "N/A"
            else:
                row_cells[3].text = "-"
        
        self.doc.add_paragraph()  # Spacing
    
    def add_bottleneck_analysis(self, 
                               bottlenecks: List[Dict[str, Any]]):
        """
        Add bottleneck analysis section.
        
        Args:
            bottlenecks: List of identified bottlenecks
        """
        self.doc.add_heading('Bottleneck Analysis', level=1)
        
        if not bottlenecks:
            success_para = self.doc.add_paragraph()
            success_run = success_para.add_run(
                "✓ No critical bottlenecks identified. Performance is within expected parameters."
            )
            success_run.font.color.rgb = ReportStyle.SUCCESS
            success_run.font.bold = True
            return
        
        # Sort by severity
        severity_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
        sorted_bottlenecks = sorted(
            bottlenecks,
            key=lambda b: (severity_order.get(b.get('severity', 'low'), 99), 
                          -b.get('impact', 0))
        )
        
        for i, bottleneck in enumerate(sorted_bottlenecks, 1):
            # Bottleneck heading
            severity = bottleneck.get('severity', 'unknown').upper()
            component = bottleneck.get('component', 'Unknown')
            
            heading = self.doc.add_heading(f"{i}. {component}", level=2)
            
            # Add severity badge
            severity_para = self.doc.add_paragraph()
            severity_run = severity_para.add_run(f"Severity: {severity}")
            severity_run.font.bold = True
            
            # Color code by severity
            if severity == 'CRITICAL':
                severity_run.font.color.rgb = ReportStyle.DANGER
            elif severity == 'HIGH':
                severity_run.font.color.rgb = ReportStyle.ACCENT
            else:
                severity_run.font.color.rgb = ReportStyle.LIGHT_GRAY
            
            # Description
            description = bottleneck.get('description', '')
            self.doc.add_paragraph(description)
            
            # Impact
            impact = bottleneck.get('impact', 0)
            impact_para = self.doc.add_paragraph()
            impact_para.add_run('Estimated Performance Impact: ')
            impact_run = impact_para.add_run(f"{impact:.1f}%")
            impact_run.font.bold = True
            impact_run.font.color.rgb = ReportStyle.ACCENT
            
            # Recommendations
            recommendations = bottleneck.get('recommendations', [])
            if recommendations:
                self.doc.add_paragraph('Recommendations:', style='List Bullet')
                for rec in recommendations:
                    self.doc.add_paragraph(rec, style='List Bullet 2')
            
            self.doc.add_paragraph()  # Spacing
    
    def add_recommendations(self, recommendations: List[str]):
        """
        Add recommendations and next steps section.
        
        Args:
            recommendations: List of recommendations
        """
        self.doc.add_heading('Recommendations & Next Steps', level=1)
        
        if not recommendations:
            self.doc.add_paragraph("Continue monitoring performance metrics.")
            return
        
        intro_para = self.doc.add_paragraph(
            "Based on the analysis, the following actions are recommended to "
            "optimize performance and address identified bottlenecks:"
        )
        
        self.doc.add_paragraph()
        
        # Numbered list of recommendations
        for i, recommendation in enumerate(recommendations, 1):
            rec_para = self.doc.add_paragraph(style='List Number')
            rec_para.add_run(recommendation)
    
    def add_technical_details(self,
                            configuration: Dict[str, Any],
                            environment: Dict[str, Any]):
        """
        Add technical specifications section.
        
        Args:
            configuration: Test configuration details
            environment: Environment/hardware details
        """
        self.doc.add_heading('Technical Specifications', level=1)
        
        # Configuration
        if configuration:
            self.doc.add_heading('Test Configuration', level=2)
            for key, value in configuration.items():
                para = self.doc.add_paragraph(style='List Bullet')
                key_run = para.add_run(f"{key}: ")
                key_run.font.bold = True
                para.add_run(str(value))
        
        # Environment
        if environment:
            self.doc.add_heading('Environment', level=2)
            for key, value in environment.items():
                para = self.doc.add_paragraph(style='List Bullet')
                key_run = para.add_run(f"{key}: ")
                key_run.font.bold = True
                para.add_run(str(value))
    
    def add_chart(self, image_path: str, caption: Optional[str] = None):
        """
        Add chart/visualization to report.
        
        Args:
            image_path: Path to image file
            caption: Optional caption
        """
        if not os.path.exists(image_path):
            return
        
        # Add image
        try:
            self.doc.add_picture(image_path, width=Inches(6.0))
            
            # Add caption
            if caption:
                caption_para = self.doc.add_paragraph(caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.runs[0]
                caption_run.font.size = Pt(ReportStyle.CAPTION_SIZE)
                caption_run.font.italic = True
                caption_run.font.color.rgb = ReportStyle.LIGHT_GRAY
        except Exception as e:
            # If image fails to load, add note
            self.doc.add_paragraph(f"[Chart: {caption or 'Visualization'}]")
    
    def generate(self, 
                data: Dict[str, Any],
                output_path: Optional[str] = None) -> str:
        """
        Generate complete report from analysis data.
        
        Args:
            data: Analysis data including results, bottlenecks, recommendations
            output_path: Where to save the report
        
        Returns:
            Path to generated report
        """
        # Cover page
        title = data.get('title', 'Performance Analysis Report')
        subtitle = data.get('subtitle', 'Technical Benchmark Results')
        company = data.get('company', 'SAAAM LLC')
        
        self.add_cover_page(title, subtitle, company)
        
        # Executive summary
        summary = data.get('summary', 'Performance analysis complete.')
        key_findings = data.get('key_findings', [])
        key_metrics = data.get('key_metrics', {})
        
        self.add_executive_summary(summary, key_findings, key_metrics)
        
        # Performance overview
        results = data.get('results', [])
        baseline = data.get('baseline')
        
        if results or self.report_type in ['comprehensive', 'benchmark']:
            self.add_performance_overview(results, baseline)
        
        # Charts
        charts = data.get('charts', {})
        for chart_type, chart_path in charts.items():
            if os.path.exists(chart_path):
                self.doc.add_page_break()
                self.doc.add_heading(f'{chart_type.replace("_", " ").title()}', level=1)
                self.add_chart(chart_path, f"Figure: {chart_type.replace('_', ' ').title()}")
        
        # Bottleneck analysis
        bottlenecks = data.get('bottlenecks', [])
        if bottlenecks or self.report_type in ['comprehensive', 'technical']:
            self.doc.add_page_break()
            self.add_bottleneck_analysis(bottlenecks)
        
        # Recommendations
        recommendations = data.get('recommendations', [])
        if recommendations or self.report_type in ['comprehensive', 'technical']:
            self.doc.add_page_break()
            self.add_recommendations(recommendations)
        
        # Technical details
        config = data.get('configuration', {})
        environment = data.get('environment', {})
        if (config or environment) and self.report_type in ['comprehensive', 'technical']:
            self.doc.add_page_break()
            self.add_technical_details(config, environment)
        
        # Save document
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            report_name = title.lower().replace(' ', '_')
            output_path = f'/mnt/user-data/outputs/{report_name}_{timestamp}.docx'
        
        self.doc.save(output_path)
        
        return output_path


def generate_report(data: Dict[str, Any], 
                   report_type: str = 'comprehensive',
                   output_path: Optional[str] = None) -> str:
    """
    Quick function to generate a technical report.
    
    Args:
        data: Analysis data
        report_type: Type of report ('executive', 'technical', 'benchmark', 'comprehensive')
        output_path: Where to save the report
    
    Returns:
        Path to generated report
    """
    generator = TechnicalReportGenerator(report_type)
    return generator.generate(data, output_path)
